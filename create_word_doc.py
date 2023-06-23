import pandas as pd
from docx import Document
from datetime import datetime
from docx.shared import Inches


def find_replace(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        paragraph.text = paragraph.text.replace(paragraph_keyword, draft_keyword)


def create_doc(companyName, sector, industry, current_price,
                fiftyTwoWeek, targetMeanPrice,
               marketCap, beta, dividendRate, companyInfo,
               shortRatio, shortPercentage, dividendHistory, analystdf,
               news, competition_df,plot,author):
    document = Document('Empty_koopvoorstel.docx')



    # replace the text in paragraphs
    for paragraph in document.paragraphs:
        paragraphtext = paragraph.text
        find_replace("[COMPANY_INFO]",companyInfo, paragraph)
        find_replace("[Short Ratio:]", shortRatio, paragraph)
        find_replace("[Short % of Shares Outstanding:]", shortPercentage, paragraph)

    ## add logo
    #logokeyword = "[Logo]"
    #for paragraph in document.paragraphs:
    #    if logokeyword in paragraph.text:
    #        r = paragraph.add_run()
    #        r.add_picture(imagelogo)


    # replace the text in tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_replace("Company:", companyName, paragraph )
                    find_replace("Sector:", sector, paragraph)
                    find_replace("Industry:", industry, paragraph)
                    find_replace("Current price:", current_price, paragraph)
                    find_replace("52-week l/h:", fiftyTwoWeek, paragraph)
                    find_replace("1-year target:", targetMeanPrice, paragraph)
                    find_replace("Market cap:", marketCap, paragraph)
                    find_replace("Beta:", beta, paragraph)
                    find_replace("Forward Dividend:", dividendRate, paragraph)
                    find_replace("Date:", "Date: " + str(datetime.today().date()), paragraph)
                    find_replace("Name:", "Name: " + str(author), paragraph)

    # input table for competitor
    competitor_df = competition_df
    table = document.add_table(competitor_df.shape[0]+1, competitor_df.shape[1])
    for j in range(competitor_df.shape[-1]):
        table.cell(0, j).text = competitor_df.columns[j]

    for i in range(competitor_df.shape[0]):
        for j in range(competitor_df.shape[-1]):
            table.cell(i + 1, j).text = str(competitor_df.values[i, j])

    # replace new table with old table
    old_competitor_table = document.tables[1]
    newTable = document.tables[4]
    old_competitor_table._element.getparent().replace(old_competitor_table._element,newTable._element )


    ## Analyst recommendations table
    analyst_df = analystdf
    table1 = document.add_table(analyst_df.shape[0]+1, analyst_df.shape[1])
    for j in range(analyst_df.shape[-1]):
        table1.cell(0,j).text = analyst_df.columns[j]

    for i in range(analyst_df.shape[0]):
        for j in range(analyst_df.shape[-1]):
            table1.cell(i+1, j).text = str(analyst_df.values[i,j])

    #replace old table with new table
    old_analyst_table = document.tables[2]
    newTable = document.tables[4]
    old_analyst_table._element.getparent().replace(old_analyst_table._element, newTable._element)
                   
    ## Dividends table

    dividend_df = dividendHistory
    dividend_df = dividend_df.reset_index()
    table2 = document.add_table(dividend_df.shape[0] + 1, dividend_df.shape[1])
    for j in range(dividend_df.shape[-1]):
        table2.cell(0, j).text = dividend_df.columns[j]

    for i in range(dividend_df.shape[0]):
        for j in range(dividend_df.shape[-1]):
            table2.cell(i + 1, j).text = str(dividend_df.values[i, j])
    # replace new table with old table
    old_dividend_table = document.tables[3]
    newTable = document.tables[4]
    old_dividend_table._element.getparent().replace(old_dividend_table._element, newTable._element)

    ## add graph
    keyword = "Price graph"
    for paragraph in document.paragraphs:
        if keyword in paragraph.text:
            r = paragraph.add_run()
            r.add_picture(plot)


    return(document)

